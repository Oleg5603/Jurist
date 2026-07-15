import io
import os
import secrets
import time
import uuid
from collections import deque

import docx
import pypdf
from fastapi import FastAPI, Request, HTTPException, Form, Depends, UploadFile, File
from fastapi.responses import RedirectResponse, FileResponse, PlainTextResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from pydantic import BaseModel

import storage
from config import APP_LOGIN, APP_PASSWORD, SESSION_SECRET
from llm import call_llm, LLMError, DEFAULT_MODEL, classify_practice
from practices import with_practice

app = FastAPI(title="Jurist")

STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")
COOKIE_NAME = "jurist_session"
COOKIE_MAX_AGE = 60 * 60 * 24 * 30  # 30 дней

_serializer = URLSafeTimedSerializer(SESSION_SECRET or "insecure-dev-secret")


def require_session(request: Request) -> None:
    token = request.cookies.get(COOKIE_NAME)
    if not token:
        raise HTTPException(status_code=401, detail="Не авторизован")
    try:
        _serializer.loads(token, max_age=COOKIE_MAX_AGE)
    except (BadSignature, SignatureExpired):
        raise HTTPException(status_code=401, detail="Сессия истекла или недействительна")


# --- Простой rate-limiting на дорогие (LLM-вызывающие) эндпоинты ---
# Единый логин/пароль на всех, без per-user лимитов: при утечке cookie
# можно быстро забить квоту OpenRouter. In-memory скользящее окно по
# session-cookie достаточно для однопроцессного hobby-scale приложения
# (сбрасывается при рестарте — не проблема).
RATE_LIMIT_MAX_REQUESTS = 20
RATE_LIMIT_WINDOW_SECONDS = 60
_rate_limit_buckets: dict[str, deque] = {}


def rate_limit(request: Request) -> None:
    token = request.cookies.get(COOKIE_NAME, "anonymous")
    now = time.monotonic()
    bucket = _rate_limit_buckets.setdefault(token, deque())
    while bucket and now - bucket[0] > RATE_LIMIT_WINDOW_SECONDS:
        bucket.popleft()
    if len(bucket) >= RATE_LIMIT_MAX_REQUESTS:
        raise HTTPException(
            status_code=429,
            detail=f"Слишком много запросов, подождите немного (лимит {RATE_LIMIT_MAX_REQUESTS} запросов в {RATE_LIMIT_WINDOW_SECONDS} секунд)",
        )
    bucket.append(now)


@app.post("/login")
def login(login: str = Form(...), password: str = Form(...)):
    login_ok = secrets.compare_digest(login, APP_LOGIN)
    password_ok = secrets.compare_digest(password, APP_PASSWORD)
    if not (login_ok and password_ok):
        raise HTTPException(status_code=401, detail="Неверный логин или пароль")
    token = _serializer.dumps({"login": login})
    response = RedirectResponse(url="/app", status_code=303)
    # Приложение торчит наружу только через HTTPS-туннели (cloudflared/ngrok),
    # поэтому Secure безопасно всегда — локальная разработка тоже обычно идёт
    # через туннель для реального теста логина.
    response.set_cookie(COOKIE_NAME, token, max_age=COOKIE_MAX_AGE, httponly=True, samesite="lax", secure=True)
    return response


@app.get("/logout")
def logout():
    response = RedirectResponse(url="/", status_code=303)
    response.delete_cookie(COOKIE_NAME)
    return response


@app.get("/")
def root():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))


@app.get("/app")
def app_page(_: None = Depends(require_session)):
    return FileResponse(os.path.join(STATIC_DIR, "app.html"))


class CaseRequest(BaseModel):
    name: str


@app.post("/api/cases")
def create_case(req: CaseRequest, _: None = Depends(require_session)):
    case_id = uuid.uuid4().hex
    return storage.save_case(case_id, req.name)


@app.get("/api/cases")
def cases_list(_: None = Depends(require_session)):
    return {"cases": storage.list_cases()}


CHAT_SYSTEM_PROMPT = """Ты — юридический ассистент, ориентированный на право Российской Федерации.
Помогаешь разобраться в правовых вопросах: даёшь ссылки на применимые нормы, объясняешь порядок действий,
указываешь на риски. ВАЖНО: твой ответ не является юридической консультацией и не заменяет очную консультацию
юриста — всегда явно указывай это, если вопрос касается конкретной спорной ситуации."""


class ChatRequest(BaseModel):
    session_id: str
    message: str
    model: str | None = None
    case_id: str | None = None


@app.post("/api/chat")
def chat(req: ChatRequest, _: None = Depends(require_session), __: None = Depends(rate_limit)):
    history = storage.load_chat(req.session_id)
    history.append({"role": "user", "content": req.message})
    practice_id = classify_practice(req.message)
    system = with_practice(CHAT_SYSTEM_PROMPT, practice_id)
    try:
        reply = call_llm(req.model or DEFAULT_MODEL, system, history)
    except LLMError as e:
        raise HTTPException(status_code=502, detail=str(e))
    history.append({"role": "assistant", "content": reply})
    storage.save_chat(req.session_id, history, case_id=req.case_id)
    return {"reply": reply}


@app.get("/api/chat/{session_id}")
def get_chat_history(session_id: str, _: None = Depends(require_session)):
    return {"messages": storage.load_chat(session_id)}


@app.get("/api/chats")
def chats_list(_: None = Depends(require_session)):
    return {"chats": storage.list_chats()}


DOCUMENT_SYSTEM_PROMPT = """Ты — юридический ассистент, готовящий черновики документов по праву РФ.
На основе вводных данных составь полный текст документа указанного типа, с корректной структурой
(шапка, стороны, предмет, условия, реквизиты для подписи). Незаполненные детали помечай [В КВАДРАТНЫХ СКОБКАХ].
В конце документа добавь пометку: "Документ сформирован автоматически, требует проверки юристом перед использованием."""


class DocumentRequest(BaseModel):
    doc_type: str
    parties: str
    subject: str
    amounts: str = ""
    extra: str = ""
    model: str | None = None
    case_id: str | None = None


@app.post("/api/documents/generate")
def generate_document(req: DocumentRequest, _: None = Depends(require_session), __: None = Depends(rate_limit)):
    user_prompt = (
        f"Тип документа: {req.doc_type}\n"
        f"Стороны: {req.parties}\n"
        f"Предмет: {req.subject}\n"
        f"Суммы: {req.amounts or 'не указаны'}\n"
        f"Доп. условия: {req.extra or 'нет'}"
    )
    practice_id = classify_practice(user_prompt)
    system = with_practice(DOCUMENT_SYSTEM_PROMPT, practice_id)
    try:
        text = call_llm(req.model or DEFAULT_MODEL, system, [{"role": "user", "content": user_prompt}])
    except LLMError as e:
        raise HTTPException(status_code=502, detail=str(e))
    doc_id = uuid.uuid4().hex
    storage.save_document(doc_id, {"doc_type": req.doc_type, "input": req.model_dump(), "text": text, "case_id": req.case_id})
    return {"id": doc_id, "text": text}


@app.get("/api/documents")
def documents_list(_: None = Depends(require_session)):
    return {"documents": storage.list_documents()}


@app.get("/api/documents/{doc_id}/download")
def download_document(doc_id: str, _: None = Depends(require_session)):
    doc = storage.load_document(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Документ не найден")
    return PlainTextResponse(
        doc["text"],
        headers={"Content-Disposition": f'attachment; filename="{doc_id}.txt"'},
    )


@app.get("/api/documents/{doc_id}/download.docx")
def download_document_docx(doc_id: str, _: None = Depends(require_session)):
    doc = storage.load_document(doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Документ не найден")
    document = docx.Document()
    for line in doc["text"].split("\n"):
        document.add_paragraph(line)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc_id}.docx"'},
    )


CONTRACT_SYSTEM_PROMPT = """Ты — юридический ассистент, анализирующий договоры по праву РФ на риски.
Изучи текст договора и составь список рисков. Каждый риск помечай одним из значков:
🔴 критический (может привести к прямым убыткам/недействительности), 🟡 важный (требует внимания),
🟢 приемлемый (незначительный, для полноты картины). Также перечисли типовые защитные пункты, которых
не хватает в договоре, и предложи конкретные формулировки правок. Отвечай структурированным списком."""


# Rough char->token safety margin below OpenRouter's 200k-token context limit,
# leaving room for the system prompt and the model's own output tokens.
MAX_CONTRACT_CHARS = 400_000


def _extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    if lower.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    if lower.endswith(".pdf"):
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise HTTPException(status_code=400, detail="Поддерживаются только файлы .txt, .docx, .pdf")


@app.post("/api/contracts/analyze")
async def analyze_contract(
    file: UploadFile = File(...),
    model: str = Form(default=""),
    case_id: str = Form(default=""),
    _: None = Depends(require_session),
    __: None = Depends(rate_limit),
):
    content = await file.read()
    text = _extract_text(file.filename, content)
    if not text.strip():
        raise HTTPException(status_code=400, detail="Не удалось извлечь текст из файла")
    if len(text) > MAX_CONTRACT_CHARS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Файл слишком большой для анализа ({len(text):,} символов, лимит {MAX_CONTRACT_CHARS:,}). "
                "Похоже, это не договор, а другой документ, либо его нужно разбить на части."
            ),
        )
    practice_id = classify_practice(text)
    system = with_practice(CONTRACT_SYSTEM_PROMPT, practice_id)
    try:
        analysis_text = call_llm(model or DEFAULT_MODEL, system, [{"role": "user", "content": text}])
    except LLMError as e:
        raise HTTPException(status_code=502, detail=str(e))
    contract_id = uuid.uuid4().hex
    storage.save_contract(contract_id, file.filename, content, {"analysis": analysis_text}, case_id=case_id or None)
    return {"id": contract_id, "analysis": analysis_text}


@app.get("/api/contracts")
def contracts_list(_: None = Depends(require_session)):
    return {"contracts": storage.list_contracts()}


@app.get("/api/contracts/{contract_id}")
def contract_detail(contract_id: str, _: None = Depends(require_session)):
    analysis = storage.load_contract_analysis(contract_id)
    if analysis is None:
        raise HTTPException(status_code=404, detail="Анализ не найден")
    return analysis


app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
